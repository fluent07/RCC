from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor

def create_report(doc_name, image_path, formulas_with_calculations):
    """
    Creates a Word document with formulas followed by their respective calculations.
    """

    # Create a new Word document
    doc = Document()

    # Adjust page margins to ensure the border is only 5mm from the edges
    section = doc.sections[0]
    section.top_margin = Inches(0.2)  # 5mm
    section.bottom_margin = Inches(0.2)  # 5mm
    section.left_margin = Inches(0.2)  # 5mm
    section.right_margin = Inches(0.2)  # 5mm

    # Add a page border (frame) close to the edges (5mm distance)
    border_xml = '''
        <w:sectPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:pgBorders w:offsetFrom="page">
                <w:top w:val="single" w:sz="12" w:space="4" w:color="000000"/>
                <w:left w:val="single" w:sz="12" w:space="4" w:color="000000"/>
                <w:bottom w:val="single" w:sz="12" w:space="4" w:color="000000"/>
                <w:right w:val="single" w:sz="12" w:space="4" w:color="000000"/>
            </w:pgBorders>
        </w:sectPr>
    '''
    section._sectPr.append(parse_xml(border_xml))

    # Add title
    title = doc.add_paragraph()
    title_run = title.add_run("Structural Analysis Report")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("\n")  # Spacer

    # Add formulas followed by calculations
    doc.add_paragraph("**Formulas & Calculations:**").bold = True

    for formula, calculation in formulas_with_calculations:
        # Add formula
        para_formula = doc.add_paragraph()
        run_formula = para_formula.add_run(formula)
        run_formula.font.name = "Courier New"  # Monospace font
        run_formula.font.size = Pt(12)
        para_formula.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add corresponding calculation below the formula
        para_calc = doc.add_paragraph()
        run_calc = para_calc.add_run(f"    → {calculation}")  # Indented
        run_calc.font.name = "Courier New"
        run_calc.font.size = Pt(12)
        run_calc.bold = True
        para_calc.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph("\n")  # Spacer

    # Insert the image at the top-right
    doc.add_paragraph("**Reinforced Concrete Section:**").bold = True
    para_img = doc.add_paragraph()
    run = para_img.add_run()
    run.add_picture(image_path, width=Inches(2.5))  # Adjust image size
    para_img.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Save the document
    doc.save(doc_name)
    print(f"Report saved as {doc_name}")

# Example usage:
formulas_with_calculations = [
    # (Formula, Calculation Result)
    ("x = (b * d^2) / (3 * E_c + A_s * E_s)", "x = (300 * 450^2) / (3 * 25E3 + 314.16 * 200E3) = 82.5 mm"),
    ("I_cr = (b * x^3) / 3 + A_s * d^2", "I_cr = (300 * 82.5^3) / 3 + 314.16 * 450^2 = 2.45E6 mm⁴"),
    ("I_uncr = (b * d^3) / 12 + A_s * d^2", "I_uncr = (300 * 450^3) / 12 + 314.16 * 450^2 = 6.88E6 mm⁴"),
]

# Run the function
create_report("structural_report_with_formulas.docx", "rc_section.png", formulas_with_calculations)
