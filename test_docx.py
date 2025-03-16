import matplotlib
matplotlib.use('Agg')  # Use non-GUI backend to avoid Tkinter issues
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import matplotlib.path as mpath
import numpy as np
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor

def create_report(doc_name, image_path, calculations):
    # Create a new Word document
    doc = Document()

    # Adjust page margins to ensure the border is only 5mm from the edges
    section = doc.sections[0]
    section.top_margin = Inches(0.2)  # 5mm
    section.bottom_margin = Inches(0.2)  # 5mm
    section.left_margin = Inches(0.2)  # 5mm
    section.right_margin = Inches(0.2)  # 5mm

    # Add a page border (frame) close to the edges (5mm distance)
    border_xml = '''
        <w:sectPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:pgBorders w:offsetFrom="page">
                <w:top w:val="single" w:sz="12" w:space="4" w:color="000000"/>
                <w:left w:val="single" w:sz="12" w:space="4" w:color="000000"/>
                <w:bottom w:val="single" w:sz="12" w:space="4" w:color="000000"/>
                <w:right w:val="single" w:sz="12" w:space="4" w:color="000000"/>
            </w:pgBorders>
        </w:sectPr>
    '''
    section._sectPr.append(parse_xml(border_xml))

    # Add title
    title = doc.add_paragraph()
    title_run = title.add_run("Structural Analysis Report")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = 1  # Center alignment

    doc.add_paragraph("\n")  # Spacer

    # Create a two-column table (Calculations on the left, Image on the right)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(4)  # Left column for calculations
    table.columns[1].width = Inches(2.5)  # Right column for image

    # Insert calculations in the left column with monospace font
    cell_left = table.cell(0, 0)
    for calc in calculations:
        para = cell_left.add_paragraph()
        run = para.add_run(f"• {calc}")
        run.font.name = "Courier New"  # Monospace font
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
        para.alignment = 0  # Left alignment

    # Insert the image in the right column
    cell_right = table.cell(0, 1)
    para_img = cell_right.add_paragraph()
    run = para_img.add_run()
    run.add_picture(image_path, width=Inches(2.5))  # Adjust image size
    para_img.alignment = 1  # Center alignment

    # Save the document
    doc.save(doc_name)
    print(f"Report saved as {doc_name}")

def draw_dimension(ax, start, end, text, vertical=False):
    """
    Draws a dimension line similar to AutoCAD.
    """
    arrow_length = 10  # Length of the arrowheads

    if vertical:
        ax.plot([start, start], [end[0], end[1]], 'k-', linewidth=1)
        ax.plot([start - arrow_length, start + arrow_length], [end[0], end[0]], 'k-', linewidth=1)
        ax.plot([start - arrow_length, start + arrow_length], [end[1], end[1]], 'k-', linewidth=1)
        ax.text(start - 15, (end[0] + end[1]) / 2, text, fontsize=9, ha='center', va='center', rotation=90, fontname="Courier New")
    else:
        ax.plot([end[0], end[1]], [start, start], 'k-', linewidth=1)
        ax.plot([end[0], end[0]], [start - arrow_length, start + arrow_length], 'k-', linewidth=1)
        ax.plot([end[1], end[1]], [start - arrow_length, start + arrow_length], 'k-', linewidth=1)
        ax.text((end[0] + end[1]) / 2, start - 15, text, fontsize=9, ha='center', va='center',fontname="Courier New")

def draw_rc_section(width, height, cover, top_bar_dia, bottom_bar_dia, stirrup_dia, top_rebars, bottom_rebars):
    """
    Draws a reinforced concrete section with AutoCAD-style dimension lines and stirrup representation.
    """

    fig, ax = plt.subplots(figsize=(6, 8))
    
    # Draw the concrete section (beam)
    rect = patches.Rectangle((0, 0), width, height, linewidth=2, edgecolor='black', facecolor='lightgray', zorder=1)
    ax.add_patch(rect)

    # Stirrup positioning
    stirrup_offset = cover - stirrup_dia  # Offset stirrups inside the cover
    inner_width = width - 2 * stirrup_offset
    inner_height = height - 2 * stirrup_offset
    bend_radius = stirrup_dia  # Radius of rounded corners

    # Draw stirrups with rounded corners using Path
    Path = mpath.Path
    path_data = [
        (Path.MOVETO, (stirrup_offset + bend_radius, stirrup_offset)),  
        (Path.LINETO, (stirrup_offset + inner_width - bend_radius, stirrup_offset)),  
        (Path.CURVE3, (stirrup_offset + inner_width, stirrup_offset)),  
        (Path.CURVE3, (stirrup_offset + inner_width, stirrup_offset + bend_radius)),  
        (Path.LINETO, (stirrup_offset + inner_width, stirrup_offset + inner_height - bend_radius)),  
        (Path.CURVE3, (stirrup_offset + inner_width, stirrup_offset + inner_height)),  
        (Path.CURVE3, (stirrup_offset + inner_width - bend_radius, stirrup_offset + inner_height)),  
        (Path.LINETO, (stirrup_offset + bend_radius, stirrup_offset + inner_height)),  
        (Path.CURVE3, (stirrup_offset, stirrup_offset + inner_height)),  
        (Path.CURVE3, (stirrup_offset, stirrup_offset + inner_height - bend_radius)),  
        (Path.LINETO, (stirrup_offset, stirrup_offset + bend_radius)),  
        (Path.CURVE3, (stirrup_offset, stirrup_offset)),  
        (Path.CURVE3, (stirrup_offset + bend_radius, stirrup_offset)),  
        (Path.CLOSEPOLY, (stirrup_offset + bend_radius, stirrup_offset))  
    ]
    
    codes, verts = zip(*path_data)
    path = mpath.Path(verts, codes)
    stirrup_patch = patches.PathPatch(path, facecolor='none', edgecolor='black', linewidth=1, zorder=2)
    ax.add_patch(stirrup_patch)

    # Reinforcement placement parameters
    rebar_spacing_top = (width - 2 * cover) / (top_rebars - 1) if top_rebars > 1 else 0
    rebar_spacing_bottom = (width - 2 * cover) / (bottom_rebars - 1) if bottom_rebars > 1 else 0
    
    # Plot top reinforcement (Compression bars - CR)
    for i in range(top_rebars):
        x = cover + i * rebar_spacing_top if top_rebars > 1 else width / 2
        y = height - cover
        ax.add_patch(plt.Circle((x, y), top_bar_dia / 2, color='blue', zorder=3))

    # Plot bottom reinforcement (Tension bars - TR)
    for i in range(bottom_rebars):
        x = cover + i * rebar_spacing_bottom if bottom_rebars > 1 else width / 2
        y = cover
        ax.add_patch(plt.Circle((x, y), bottom_bar_dia / 2, color='red', zorder=3))

    # Draw AutoCAD-style dimension lines
    draw_dimension(ax, -30, (0, height), text=f"{height} mm", vertical=True)  # Height
    draw_dimension(ax, height + 30, (0, width), text=f"{width} mm", vertical=False)  # Width

    # Labels
    ax.text(width + 20, height - cover, f'CR: {top_rebars} x {top_bar_dia} mm', fontsize=14, va='center', fontname="Courier New")
    ax.text(width + 20, cover, f'TR: {bottom_rebars} x {bottom_bar_dia} mm', fontsize=14, va='center',fontname="Courier New")
    ax.text(width + 20, height / 2, f'SR: {stirrup_dia} mm', fontsize=14, va='center',fontname="Courier New")

    # Formatting
    ax.set_xlim(-50, width + 50)
    ax.set_ylim(-50, height + 50)
    ax.set_aspect('equal')
    ax.axis('off')

    # Save the image
    plt.savefig("rc_section.png", dpi=300, bbox_inches='tight')
    print("Image saved as rc_section.png")


# Example usage:
draw_rc_section(
    width=300, 
    height=450, 
    cover=30, 
    top_bar_dia=14,   
    bottom_bar_dia=20, 
    stirrup_dia=8,     
    top_rebars=2, 
    bottom_rebars=4
)

# Example usage:
calculations = [
    "Moment capacity: 45 kNm",
    "Neutral axis depth (x/d): 0.38",
    "Shear reinforcement: ϕ8 @ 150 mm",
    "Deflection check: OK",
]

# Run the function
create_report("structural_report.docx", "rc_section.png", calculations)