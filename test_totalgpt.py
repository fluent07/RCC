import matplotlib
matplotlib.use('Agg')  # Use non-GUI backend to avoid Tkinter issues
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import math as m
import matplotlib.path as mpath
import numpy as np
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Mm, Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

# --- Given Data ---
b = 300  # mm, width of section
h = 450  # mm, total height
c = 40  # mm, cover from top and bottom
d = h - c
n = 15  # Modular ratio
tr = 20
n_tr = 4
cr = 14
n_cr =2
# Reinforcement details
Asc = n_cr * (m.pi / 4) * cr**2  # Compression reinforcement area (mm^2)
As = n_tr * (m.pi / 4) * tr**2  # Tension reinforcement area (mm^2)

# Material properties
f_ck = 25  # MPa, characteristic compressive strength of concrete
f_y = 450  # MPa, yield strength of steel
gamma_c = 1.5  # Safety factor for concrete
gamma_s = 1.15  # Safety factor for steel
alpha_cc = 0.85  # Concrete strength reduction factor
Es = 210000  # Young's modulus of steel

# Design strengths
fcd = alpha_cc * (f_ck / gamma_c)  # MPa
fyd = f_y / gamma_s  # MPa

# Function to calculate neutral axis using Newton-Raphson
def newtonR(f, x0):
    tolf = 1e-08
    tolx = 1e-08
    dx = 1e-08
    df = lambda x: (f(x + dx) - f(x)) / dx
    iter_root = [x0]
    for k in range(1000):
        iter_root.append(iter_root[k] - (f(iter_root[k]) / df(iter_root[k])))
        if abs(f(iter_root[k + 1])) <= tolf or abs(iter_root[k + 1] - iter_root[k]) <= tolx:
            return iter_root[k + 1]
    raise ValueError("Newton-Raphson method did not converge.")

# --- Calculations ---
# Ultimate Neutral axis depth (x) from force equilibrium
x_ult = (As - Asc) * fyd / (0.8 * fcd * b)

# Design moment resistance (Mrd)
M_rd = As * fyd * (d - 0.4 * x_ult) / (10**6)  # Convert to kNm
M_rd2 = (fcd * 0.8 * x_ult * b * (d - 0.4 * x_ult) + fyd * Asc * (d - c)) * 1e-6

# Uncracked moment of inertia (I_uncr) - for a rectangular section
Eqn_NA_Uncracked = (
    lambda x_uncr: ((b * x_uncr**2) / 2)
    + (n * Asc * (x_uncr - c))
    - ((b * (h - x_uncr) ** 2) / 2)
    - (n * As * (d - x_uncr))
)
NA_uncracked = newtonR(Eqn_NA_Uncracked, 100)  # Neutral axis of uncracked section
# print("NA_uncracked:",NA_uncracked,"mm")

I_uncr = (
    (b / 3) * NA_uncracked**3
    + n * Asc * (NA_uncracked - c) ** 2
    + (b / 3) * (h - NA_uncracked) ** 3
    + n * As * (d - NA_uncracked) ** 2
)

# Cracked moment of inertia (I_cr) - transformed section
Eqn_NA_cracked = lambda x_cr: (b * x_cr**2 / 2) + (n * Asc * (x_cr - c)) - (n * As * (d - x_cr))
NA_cracked = newtonR(Eqn_NA_cracked, 100)

# Parallel axis theorem to find I_cr
I_cr = (
    (b / 3) * NA_cracked**3
    + n * Asc * (NA_cracked - c) ** 2
    + n * As * (d - NA_cracked) ** 2
)  # (b * x**3) / 3 + (n * As * (d - y_n)**2) + (n * Asc * (y_n - c)**2)

# First moment of cracking (M_cr)
f_ctm = 0.3 * (f_ck ** (2 / 3))  # MPa, mean tensile strength of concrete
y = h / 2  # Distance to extreme fiber

M_cr = (f_ctm * I_uncr) / (h - NA_uncracked) / (10**6)  # Convert to kNm

results = {
    "Neutral Axis (Uncracked)": (f"{NA_uncracked:.3f} mm", "((b * x_uncr^2) / 2)+(n * Asc * (x_uncr - c)) = ((b * (h - x_uncr)^2) / 2)+(n * As * (d - x_uncr))"),
    "Moment of Inertia (Uncracked)": (I_uncr, "(b / 3) * NA_uncracked^3 + n * Asc * (NA_uncracked - c)^2 + (b / 3) * (h - NA_uncracked)^3 + n * As * (d - NA_uncracked)^2"),
    "Neutral Axis (Cracked)": (f"{NA_cracked:.3f} mm", "(b * x_cr^2 / 2) + (n * Asc * (x_cr - c)) = (n * As * (d - x_cr))"),
    "Moment of Inertia (Cracked)": (I_cr, "(b / 3) * NA_cracked^3 + n * Asc * (NA_cracked - c)^2 + n * As * (d - NA_cracked)^2"),
    "Neutral Axis Depth at failure (x_ult)": (f"{x_ult:.3f} mm", "(As - Asc) * fyd / (0.8 * fcd * b)"),
    "Moment of Resistance M_rd": (M_rd, "M_rd = As * fyd * (d - 0.4 * x)"),
    "Moment of Resistance M_rd2": (M_rd2, "M_rd2 = (fcd * 0.8 * x * b * (d - 0.4 * x) + fyd * Asc * (d - c))"),
    "Moment of First Crack (M_cr)": (M_cr, "M_cr = (f_ctm * I_uncr) / (h - x)"),
}

def draw_dimension(ax, start, end, text, vertical=False):
    """
    Draws a dimension line similar to AutoCAD.
    """
    arrow_length = 10  # Length of the arrowheads
    font_size = 14  # Font size for the dimension text

    if vertical:
        ax.plot([start, start], [end[0], end[1]], 'k-', linewidth=1)
        ax.plot([start - arrow_length, start + arrow_length], [end[0], end[0]], 'k-', linewidth=1)
        ax.plot([start - arrow_length, start + arrow_length], [end[1], end[1]], 'k-', linewidth=1)
        ax.text(start - 15, (end[0] + end[1]) / 2, text, fontsize=font_size, ha='center', va='bottom', rotation=90, fontname="Courier New")
    else:
        ax.plot([end[0], end[1]], [start, start], 'k-', linewidth=1)
        ax.plot([end[0], end[0]], [start - arrow_length, start + arrow_length], 'k-', linewidth=1)
        ax.plot([end[1], end[1]], [start - arrow_length, start + arrow_length], 'k-', linewidth=1)
        ax.text((end[0] + end[1]) / 2, start + 15, text, fontsize=font_size, ha='center', va='bottom', fontname="Courier New")

def draw_rc_section(width, height, cover, top_bar_dia, bottom_bar_dia, stirrup_dia, top_rebars, bottom_rebars):
    """
    Draws a reinforced concrete section with AutoCAD-style dimension lines and stirrup representation.
    """

    fig, ax = plt.subplots(figsize=(4, 6))
    
    # Draw the concrete section (beam)
    rect = patches.Rectangle((0, 0), width, height, linewidth=2, edgecolor='black', facecolor='lightgray', zorder=1)
    ax.add_patch(rect)

    # Stirrup positioning
    stirrup_offset = cover - stirrup_dia  # Offset stirrups inside the cover
    inner_width = width - 2 * stirrup_offset
    inner_height = height - 2 * stirrup_offset
    bend_radius = stirrup_dia  # Radius of rounded corners

    # Draw stirrups with rounded corners using Path
    Path = mpath.Path
    path_data = [
        (Path.MOVETO, (stirrup_offset + bend_radius, stirrup_offset)),  
        (Path.LINETO, (stirrup_offset + inner_width - bend_radius, stirrup_offset)),  
        (Path.CURVE3, (stirrup_offset + inner_width, stirrup_offset)),  
        (Path.CURVE3, (stirrup_offset + inner_width, stirrup_offset + bend_radius)),  
        (Path.LINETO, (stirrup_offset + inner_width, stirrup_offset + inner_height - bend_radius)),  
        (Path.CURVE3, (stirrup_offset + inner_width, stirrup_offset + inner_height)),  
        (Path.CURVE3, (stirrup_offset + inner_width - bend_radius, stirrup_offset + inner_height)),  
        (Path.LINETO, (stirrup_offset + bend_radius, stirrup_offset + inner_height)),  
        (Path.CURVE3, (stirrup_offset, stirrup_offset + inner_height)),  
        (Path.CURVE3, (stirrup_offset, stirrup_offset + inner_height - bend_radius)),  
        (Path.LINETO, (stirrup_offset, stirrup_offset + bend_radius)),  
        (Path.CURVE3, (stirrup_offset, stirrup_offset)),  
        (Path.CURVE3, (stirrup_offset + bend_radius, stirrup_offset)),  
        (Path.CLOSEPOLY, (stirrup_offset + bend_radius, stirrup_offset))  
    ]
    
    codes, verts = zip(*path_data)
    path = mpath.Path(verts, codes)
    stirrup_patch = patches.PathPatch(path, facecolor='none', edgecolor='black', linewidth=1, zorder=2)
    ax.add_patch(stirrup_patch)

    # Reinforcement placement parameters
    rebar_spacing_top = (width - 2 * cover) / (top_rebars - 1) if top_rebars > 1 else 0
    rebar_spacing_bottom = (width - 2 * cover) / (bottom_rebars - 1) if bottom_rebars > 1 else 0
    
    # Plot top reinforcement (Compression bars - CR)
    for i in range(top_rebars):
        x = cover + i * rebar_spacing_top if top_rebars > 1 else width / 2
        y = height - cover
        ax.add_patch(plt.Circle((x, y), top_bar_dia / 2, color='blue', zorder=3))

    # Plot bottom reinforcement (Tension bars - TR)
    for i in range(bottom_rebars):
        x = cover + i * rebar_spacing_bottom if bottom_rebars > 1 else width / 2
        y = cover
        ax.add_patch(plt.Circle((x, y), bottom_bar_dia / 2, color='red', zorder=3))

    # Draw AutoCAD-style dimension lines
    draw_dimension(ax, -30, (0, height), text=f"{height} mm", vertical=True)  # Height
    draw_dimension(ax, height + 30, (0, width), text=f"{width} mm", vertical=False)  # Width

    # Labels
    ax.text(width + 20, height - cover, f'CR: {top_rebars} x {top_bar_dia} mm', fontsize=14, va='center', fontname="Courier New")
    ax.text(width + 20, cover, f'TR: {bottom_rebars} x {bottom_bar_dia} mm', fontsize=14, va='center',fontname="Courier New")
    ax.text(width + 20, height / 2, f'SR: {stirrup_dia} mm', fontsize=14, va='center',fontname="Courier New")

    # Formatting
    ax.set_xlim(-50, width + 50)
    ax.set_ylim(-50, height + 50)
    ax.set_aspect('equal')
    ax.axis('off')

    # Save the image
    plt.savefig("rc_section.png", dpi=450, bbox_inches='tight')
    print("Image saved as rc_section.png")


# Example usage:
draw_rc_section(
    width=b, 
    height=h, 
    cover=c, 
    top_bar_dia=cr,   
    bottom_bar_dia=tr, 
    stirrup_dia=8,     
    top_rebars=n_cr, 
    bottom_rebars=n_tr
)

# --- Save to Word Document ---


def create_report():
    doc = Document()

    # Set the page margins to 10mm
    section = doc.sections[0]
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)
    section.left_margin = Mm(10)
    section.right_margin = Mm(10)

    # Create a border for the page
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')

    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')  # 4/8 pt = 0.5 pt
        border.set(qn('w:space'), '100')  # 100 twips = 5 mm
        border.set(qn('w:color'), '000000')
        pgBorders.append(border)

    sectPr.append(pgBorders)

    # Add title
    title = doc.add_paragraph("Section Analysis Report", style="Title")
    title.runs[0].font.name = 'Twilio Sans Mono'
    title.runs[0].font.size = Pt(12)

    # Add heading for Reinforced Concrete Section Details
    heading = doc.add_paragraph("**Reinforced Concrete Section Details:**")
    heading.runs[0].font.name = 'Twilio Sans Mono'
    heading.runs[0].font.size = Pt(10)
    heading.runs[0].bold = True

    # Create a table for the image and beam section details
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    # Set the width of the first column to 11.9 cm
    table.cell(0, 0).width = Mm(119)
    # Set the width of the second column to 8 cm
    table.cell(0, 1).width = Mm(80)

#     # Set table borders
#     tbl = table._tbl  # Get the table element
#     tblBorders = OxmlElement('w:tblBorders')
#     for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
#             border = OxmlElement(f'w:{border_name}')
#             border.set(qn('w:val'), 'single')
#             border.set(qn('w:sz'), '4')  # 4/8 pt = 0.5 pt
#             border.set(qn('w:space'), '0')
#             border.set(qn('w:color'), '000000')
#             tblBorders.append(border)
#     tbl.tblPr.append(tblBorders)

    details = [
        f"Width - b: {b} mm",
        f"Height - h: {h} mm",
        f"Cover - c: {c} mm",
        f"Effective depth - d: {h} mm - {c} mm = {h-c} mm",
        f"Modular ratio - n: {n}",
        f"Top Reinforcement - Asc: {n_cr} x {cr} mm = {round(Asc,2)} mm^2",
        f"Bottom Reinforcement - As: {n_tr} x {tr} mm = {round(As,2)} mm^2\n",
        f"Concrete Strength : fck = {f_ck} MPa",
        f"Concrete Safety Factor : γc = {gamma_c}",
        f"Concrete Strength Reduction Factor : αcc = {alpha_cc}",
        f"Factored Concrete Compressive Strength : fcd = {round(fcd,2)} MPa\n",
        f"Steel Yield Strength : fyk = {f_y} MPa",
        f"Steel Safety Factor : γs = {gamma_s}",
        f"Factored Steel Yield Strength : fyd = {round(fyd,2)} MPa"
    ]

    # Add beam section details to the left cell
    cell = table.cell(0, 0)
    details_paragraph = cell.add_paragraph()
    details_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    for detail in details:
        parts = detail.split(": ")
        run = details_paragraph.add_run(f'{parts[0]}: ')
        run.font.name = 'Twilio Sans Mono'
        run.font.size = Pt(8)
        run = details_paragraph.add_run(f'{parts[1]}\n')
        run.font.name = 'Twilio Sans Mono'
        run.font.size = Pt(8)
        run.bold = True

    # Add image to the right cell
    cell = table.cell(0, 1)
    cell_paragraph = cell.add_paragraph()
    cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = cell_paragraph.add_run()
    run.add_picture("rc_section.png", width=Inches(3))

    # Add calculations section
    heading = doc.add_paragraph("\n**Formulas & Calculations:**", style="Heading 1")
    heading.runs[0].font.name = 'Twilio Sans Mono'
    heading.runs[0].font.size = Pt(12)

    # Add all results with specific spacing
    for key, (value, formula) in results.items():
        results_paragraph = doc.add_paragraph()
        results_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        results_paragraph.paragraph_format.line_spacing = Pt(10)  # 6pt spacing between lines within the same key-value pair
        results_paragraph.paragraph_format.space_after = Pt(10)  # 10pt space after each result line

        run = results_paragraph.add_run(f"{key}:\n")
        # run.bold = True
        run.font.name = 'Twilio Sans Mono'
        run.font.size = Pt(8)
        
        # Add formula as plain text
        run = results_paragraph.add_run(f"    Formula: {formula}\n")
        run.italic = True
        run.font.name = 'Twilio Sans Mono'
        run.font.size = Pt(8)
        
        # Add result in the same paragraph, on a new line
        run = results_paragraph.add_run(f"    Result: {value}\n")
        run.bold = True
        run.font.name = 'Twilio Sans Mono'
        run.font.size = Pt(8)

    doc.save("structural_report.docx")

create_report()
print("Report saved as structural_report.docx")
